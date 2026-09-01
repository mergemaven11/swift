from docx import Document
from pypdf import PdfWriter

from swift_files.docx_ops import extract_docx, inspect_docx
from swift_files.pdf_ops import extract_pdf, inspect_pdf


def test_docx_inspect_and_extract(tmp_path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.core_properties.title = "Runbook"
    doc.add_paragraph("Platform engineering is fun")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "service"
    table.cell(0, 1).text = "healthy"
    doc.save(path)
    info = inspect_docx(path)
    assert info["title"] == "Runbook" and info["tables"] == 1
    output = tmp_path / "sample.txt"
    extract_docx(path, output)
    text = output.read_text(encoding="utf-8")
    assert "Platform engineering" in text and "service\thealthy" in text


def test_pdf_inspect_and_extract(tmp_path):
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.add_metadata({"/Title": "Artifact"})
    with path.open("wb") as handle:
        writer.write(handle)
    info = inspect_pdf(path)
    assert info["pages"] == 1 and info["title"] == "Artifact"
    output = tmp_path / "sample.txt"
    extract_pdf(path, output)
    assert output.exists()
