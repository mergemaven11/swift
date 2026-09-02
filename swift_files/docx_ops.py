"""Inspect, extract, and safely copy Microsoft Word DOCX artifacts."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .core import SwiftFilezError, atomic_write_text, safe_copy


def _document(path: str | Path) -> Document:
    """Open a DOCX file as a python-docx document.

    Args:
        path: DOCX file to open.

    Returns:
        Parsed Word document.

    Raises:
        SwiftFilezError: If the file is missing or cannot be parsed as DOCX.
    """
    docx_path = Path(path)
    if not docx_path.is_file():
        raise SwiftFilezError(f"DOCX file not found: {docx_path}")
    try:
        return Document(str(docx_path))
    except Exception as exc:
        raise SwiftFilezError(f"Could not open DOCX file: {docx_path}") from exc


def extract_text(path: str | Path) -> str:
    """Extract paragraph and table text from a DOCX file.

    Table cells are joined with tabs while document blocks are joined with
    newlines so command-line output remains readable.

    Args:
        path: DOCX file to extract.

    Returns:
        Plain text from nonempty paragraphs and all table rows.

    Raises:
        SwiftFilezError: If the document cannot be opened.
    """
    document = _document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def inspect_docx(path: str | Path) -> dict:
    """Return structural, text, and core-metadata details for a DOCX file.

    Args:
        path: DOCX file to inspect.

    Returns:
        Mapping containing paragraph/table counts, text size statistics, and
        title/author metadata when present.

    Raises:
        SwiftFilezError: If the document cannot be opened.
    """
    document = _document(path)
    text = extract_text(path)
    return {
        "path": str(Path(path)),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "words": len(text.split()),
        "characters": len(text),
        "title": document.core_properties.title or None,
        "author": document.core_properties.author or None,
    }


def extract_docx(path: str | Path, output: str | Path) -> Path:
    """Extract DOCX text and atomically write it to a text file.

    Args:
        path: Source DOCX file.
        output: Destination text file.

    Returns:
        Path to the written output file.

    Raises:
        SwiftFilezError: If the document cannot be opened.
    """
    return atomic_write_text(output, extract_text(path) + "\n")


def copy_docx(path: str | Path, output: str | Path) -> Path:
    """Safely copy a DOCX artifact without modifying its contents.

    Args:
        path: Source DOCX file.
        output: Destination path.

    Returns:
        Path to the copied file.
    """
    return safe_copy(path, output)
